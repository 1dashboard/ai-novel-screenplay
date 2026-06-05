"""Word (.docx) novel parser."""

from __future__ import annotations

import re
from pathlib import Path

from .base import BaseParser, NovelText, Chapter

try:
    import docx
except ImportError:
    docx = None


class DocxParser(BaseParser):
    """Parser for Microsoft Word (.docx) novel files."""

    CHAPTER_PATTERNS = [
        re.compile(r"^第[零一二三四五六七八九十百千\d]+章[：:\s]*(.*)", re.IGNORECASE),
        re.compile(r"^Chapter\s+\d+[：:\s]*(.*)", re.IGNORECASE),
    ]

    def parse(self, file_path: str) -> NovelText:
        if docx is None:
            raise ImportError("python-docx is required to parse .docx files. Install it: pip install python-docx")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect misnamed .doc files (OLE2 magic bytes)
        header = path.read_bytes()[:8]
        if header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
            # This is actually a legacy .doc file — delegate
            from .doc_parser import DocParser
            return DocParser().parse(file_path)

        doc = docx.Document(str(path))

        # Extract paragraphs as lines of text
        lines: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
            else:
                lines.append("")  # preserve paragraph breaks

        title = self._detect_title_from_docx(doc)
        chapters = self._split_chapters(lines)

        if len(chapters) < 3:
            from .txt_parser import TxtParser
            fallback = TxtParser()
            chapters = fallback._split_chapters(lines)

        return NovelText(
            title=title,
            chapters=chapters,
            source_path=str(path.resolve()),
        )

    def _detect_title_from_docx(self, doc) -> str:
        """Try to find a title from the first large/centered paragraph."""
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            if not text:
                continue
            # Title paragraphs are often centered or have larger font
            if para.alignment == 1:  # CENTER
                return text
            for run in para.runs:
                if run.font.size and run.font.size >= 16 * 12700:  # 16pt+
                    return text
                break
        # Fallback
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            if text and len(text) <= 60:
                return text
        return "未命名作品"

    def _split_chapters(self, lines: list[str]) -> list[Chapter]:
        """Split lines into chapters."""
        chapters: list[Chapter] = []
        current_lines: list[str] = []
        current_title = ""

        for line in lines:
            matched = False
            for pat in self.CHAPTER_PATTERNS:
                m = pat.match(line)
                if m:
                    if current_lines:
                        content = "\n".join(current_lines).strip()
                        if content:
                            chapters.append(Chapter(
                                number=len(chapters) + 1,
                                title=current_title or f"第{len(chapters) + 1}章",
                                content=content,
                            ))
                    current_title = m.group(1).strip() if m.lastindex else line
                    current_lines = []
                    matched = True
                    break
            if not matched:
                current_lines.append(line)

        if current_lines:
            content = "\n".join(current_lines).strip()
            if content:
                chapters.append(Chapter(
                    number=len(chapters) + 1,
                    title=current_title or f"第{len(chapters) + 1}章",
                    content=content,
                ))

        return chapters
